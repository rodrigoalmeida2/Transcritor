from fpdf import FPDF
from docx import Document

# salva resumos em PDF e Word e txt
class SaveSummaries:
    @staticmethod
    def save_to_pdf(text, output_file="text.pdf"):
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)

        pdf.cell(200, 10, txt="Transcrição", ln=True, align='C')
        pdf.ln(10)  # Adiciona uma linha em branco

        
        pdf.multi_cell(0, 10, txt=f"{text}", align='L')

        pdf.output(output_file)
        print(f"Resumo salvo em PDF: {output_file}")

    @staticmethod
    def save_to_word(text, output_file="text.docx"):
        """Salva os resumos em um arquivo Word."""
        document = Document()
        document.add_heading("Transcrição", level=1)
        document.add_paragraph(f"{text}")

        document.save(output_file)
        print(f"Resumo salvo em Word: {output_file}")
